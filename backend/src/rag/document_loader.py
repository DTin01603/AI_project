"""Document loading utilities for multi-format indexing."""

from __future__ import annotations

import ast
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


SUPPORTED_CODE_EXTENSIONS = {".py", ".js", ".ts", ".java"}


class DocumentLoadError(RuntimeError):
    """Raised when a document cannot be loaded."""


@dataclass
class DocumentMetadata:
    """Metadata extracted from a source file."""

    file_name: str
    file_path: str
    source_type: str
    file_size: int
    created_at: str
    modified_at: str
    file_extension: str
    extra: dict[str, Any]


@dataclass
class Document:
    """Loaded document payload used by the indexing pipeline."""

    id: str
    text: str
    metadata: DocumentMetadata
    source_type: str


class DocumentLoader(ABC):
    """Base interface for file loaders."""

    @abstractmethod
    def supports(self, file_path: Path) -> bool:
        """Return True when loader can parse the provided file."""

    @abstractmethod
    def load(self, file_path: Path) -> Document:
        """Load and parse one file into a Document."""

    def _build_metadata(self, file_path: Path, source_type: str, extra: dict[str, Any]) -> DocumentMetadata:
        stat = file_path.stat()
        return DocumentMetadata(
            file_name=file_path.name,
            file_path=str(file_path),
            source_type=source_type,
            file_size=int(stat.st_size),
            created_at=datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat(),
            modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            file_extension=file_path.suffix.lower(),
            extra=extra,
        )


class TextLoader(DocumentLoader):
    """Loader for plain text documents."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".txt", ".text", ".log", ".csv", ".rst"}

    def load(self, file_path: Path) -> Document:
        text = _read_text_with_detection(file_path)
        metadata = self._build_metadata(file_path, source_type="document", extra={"loader": "text"})
        return Document(id="", text=text, metadata=metadata, source_type="document")


class MarkdownLoader(DocumentLoader):
    """Loader for markdown files with optional frontmatter."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".md", ".markdown"}

    def load(self, file_path: Path) -> Document:
        raw = _read_text_with_detection(file_path)
        frontmatter, body = _extract_frontmatter(raw)
        metadata = self._build_metadata(
            file_path,
            source_type="document",
            extra={"loader": "markdown", "frontmatter": frontmatter},
        )
        return Document(id="", text=body, metadata=metadata, source_type="document")


class PDFLoader(DocumentLoader):
    """Loader for PDF files."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def load(self, file_path: Path) -> Document:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            try:
                from PyPDF2 import PdfReader  # type: ignore
            except Exception as exc:
                raise DocumentLoadError(
                    "PDF dependencies missing. Install 'pypdf' or 'PyPDF2'."
                ) from exc

        try:
            reader = PdfReader(str(file_path))
            pages: list[str] = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            text = "\n\n".join(pages)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to load PDF '{file_path.name}': {exc}") from exc

        metadata = self._build_metadata(
            file_path,
            source_type="document",
            extra={"loader": "pdf", "page_count": len(getattr(reader, "pages", []))},
        )
        return Document(id="", text=text, metadata=metadata, source_type="document")


class DOCXLoader(DocumentLoader):
    """Loader for DOCX files."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def load(self, file_path: Path) -> Document:
        try:
            import docx  # type: ignore
        except Exception as exc:
            raise DocumentLoadError("DOCX dependency missing. Install 'python-docx'.") from exc

        try:
            doc = docx.Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            core_props = doc.core_properties
            extra = {
                "loader": "docx",
                "author": getattr(core_props, "author", None),
                "title": getattr(core_props, "title", None),
                "subject": getattr(core_props, "subject", None),
            }
        except Exception as exc:
            raise DocumentLoadError(f"Failed to load DOCX '{file_path.name}': {exc}") from exc

        metadata = self._build_metadata(file_path, source_type="document", extra=extra)
        return Document(id="", text=text, metadata=metadata, source_type="document")


class CodeLoader(DocumentLoader):
    """Loader for source code files."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in SUPPORTED_CODE_EXTENSIONS

    def load(self, file_path: Path) -> Document:
        text = _read_text_with_detection(file_path)
        ext = file_path.suffix.lower()

        signatures: list[str] = []
        comments: list[str] = []
        if ext == ".py":
            signatures, comments = _extract_python_features(text)
        else:
            signatures, comments = _extract_generic_code_features(text)

        metadata = self._build_metadata(
            file_path,
            source_type="code_file",
            extra={
                "loader": "code",
                "language": ext.lstrip("."),
                "signatures": signatures,
                "comments": comments,
            },
        )
        return Document(id="", text=text, metadata=metadata, source_type="code_file")


def get_default_loaders() -> list[DocumentLoader]:
    """Return default loader chain ordered by specificity."""
    return [PDFLoader(), DOCXLoader(), MarkdownLoader(), CodeLoader(), TextLoader()]


def load_document(file_path: str | Path, loaders: list[DocumentLoader] | None = None) -> Document:
    """Load one document using the first supporting loader."""
    resolved = Path(file_path)
    if not resolved.exists() or not resolved.is_file():
        raise DocumentLoadError(f"File does not exist: {resolved}")

    loader_chain = loaders or get_default_loaders()
    for loader in loader_chain:
        if loader.supports(resolved):
            return loader.load(resolved)

    raise DocumentLoadError(f"Unsupported document format: {resolved.suffix.lower() or '<none>'}")


def _extract_frontmatter(text: str) -> tuple[dict[str, str], str]:
    normalized = text.replace("\r\n", "\n")

    if not normalized.startswith("---\n"):
        return {}, text

    end = normalized.find("\n---\n", 4)
    if end < 0:
        return {}, text

    fm_raw = normalized[4:end]
    body = normalized[end + 5 :]
    frontmatter: dict[str, str] = {}

    for line in fm_raw.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        frontmatter[key.strip()] = value.strip()

    return frontmatter, body


def _read_text_with_detection(file_path: Path) -> str:
    raw = file_path.read_bytes()

    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    try:
        import chardet  # type: ignore

        detected = chardet.detect(raw)
        enc = detected.get("encoding") or "utf-8"
        return raw.decode(enc, errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


def _extract_python_features(text: str) -> tuple[list[str], list[str]]:
    signatures: list[str] = []
    comments: list[str] = []

    try:
        tree = ast.parse(text)
        doc = ast.get_docstring(tree)
        if doc:
            comments.append(doc)

        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef):
                signatures.append(f"def {node.name}(...)")
            elif isinstance(node, ast.AsyncFunctionDef):
                signatures.append(f"async def {node.name}(...)")
            elif isinstance(node, ast.ClassDef):
                signatures.append(f"class {node.name}")
    except Exception:
        pass

    comments.extend(re.findall(r"#\s*(.+)", text))
    return signatures[:200], comments[:200]


def _extract_generic_code_features(text: str) -> tuple[list[str], list[str]]:
    signatures = re.findall(
        r"(?:class|interface|function|def)\s+[A-Za-z_][A-Za-z0-9_]*|(?:public|private|protected)\s+[A-Za-z0-9_<>,\[\]]+\s+[A-Za-z_][A-Za-z0-9_]*\s*\(",
        text,
    )
    comments = re.findall(r"//\s*(.+)", text)
    comments.extend(re.findall(r"/\*([\s\S]*?)\*/", text))
    return signatures[:200], comments[:200]
